# import aspose.words as aw
#
# # load Word document
# doc = aw.Document("N65236-AIS-TPLAN-0282.docx")
#
# # replace text
# doc.range.replace("106.0.1370.47", "107.0.1418.24", aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))
#
# # save the modified document
# doc.save("N65236-AIS-TPLAN-0282.docx")

from docx import Document
import os
from docx.shared import Pt

template_file_path = 'N65236-AIS-TPLAN-0282.docx'
output_file_path = 'N65236-AIS-TPLAN-0282A.docx'

variables = {
    "106.0.1370.47": "107.0.1418.24",
}

template_document = Document(template_file_path)
style = template_document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)


def main():
    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            #print(variable_key)
            #print(variable_value)
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        print(paragraph.text)
        if key in paragraph.text:
            print(key)
            paragraph.text = paragraph.text.replace(key, value)
            paragraph.style = template_document.styles['Normal']
            print(paragraph.text)
        # inline = paragraph.runs
        # for item in inline:
        #     #print(item.text)
        #     if key in item.text:
        #         item.text = item.text.replace(key, value)


if __name__ == '__main__':
    main()
